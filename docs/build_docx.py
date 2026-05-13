# -*- coding: utf-8 -*-
"""从同目录 用户说明.md 生成 用户说明.docx。依赖：pip install python-docx"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("请先安装: pip install python-docx", file=sys.stderr)
    sys.exit(1)

HERE = Path(__file__).resolve().parent
# 使用 Unicode 转义，避免源文件编码导致输出文件名乱码
MD_PATH = HERE / "\u7528\u6237\u8bf4\u660e.md"
OUT_PATH_ZH = HERE / "\u7528\u6237\u8bf4\u660e.docx"
OUT_PATH_ASCII = HERE / "USER_GUIDE_zh-CN.docx"


def _plain(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    return s


def _flush_buffer(doc: Document, buf: list[str]) -> None:
    if not buf:
        return
    text = "\n".join(buf).strip()
    if not text:
        return
    doc.add_paragraph(_plain(text))


def main() -> None:
    if not MD_PATH.is_file():
        raise SystemExit(f"缺少文件: {MD_PATH}")

    raw = MD_PATH.read_text(encoding="utf-8")
    lines = raw.splitlines()
    doc = Document()
    buf: list[str] = []
    in_fence = False
    fence: list[str] = []

    for line in lines:
        if line.strip().startswith("```"):
            if in_fence:
                _flush_buffer(doc, buf)
                buf = []
                p = doc.add_paragraph()
                run = p.add_run("\n".join(fence))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                fence = []
                in_fence = False
            else:
                _flush_buffer(doc, buf)
                buf = []
                in_fence = True
            continue

        if in_fence:
            fence.append(line)
            continue

        if line.strip() == "---":
            _flush_buffer(doc, buf)
            buf = []
            doc.add_paragraph("─" * 28)
            continue

        if line.startswith("### "):
            _flush_buffer(doc, buf)
            doc.add_heading(_plain(line[4:].strip()), level=2)
            continue
        if line.startswith("## "):
            _flush_buffer(doc, buf)
            doc.add_heading(_plain(line[3:].strip()), level=1)
            continue
        if line.startswith("# "):
            _flush_buffer(doc, buf)
            doc.add_heading(_plain(line[2:].strip()), level=0)
            continue

        if line.startswith("- "):
            _flush_buffer(doc, buf)
            doc.add_paragraph(_plain(line[2:].strip()), style="List Bullet")
            continue

        if not line.strip():
            _flush_buffer(doc, buf)
            buf = []
            continue

        buf.append(line)

    _flush_buffer(doc, buf)
    doc.save(OUT_PATH_ZH)
    doc.save(OUT_PATH_ASCII)
    print(f"已生成: {OUT_PATH_ZH}")
    print(f"已生成: {OUT_PATH_ASCII}")


if __name__ == "__main__":
    main()
